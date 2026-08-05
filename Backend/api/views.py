from .validators import is_cs_resume

from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status
from resume.models import Resume
from skills.models import Skill
from skills_extractor import extract_skills
# import requests
# import fitz  # PyMuPDF
# from docx import Document
# from io import BytesIO
from urllib.parse import urlparse
from PIL import Image
import pytesseract
import requests
import fitz
import textract

from docx import Document
from io import BytesIO
from pathlib import Path

def extract_text_from_pdf_bytes(file_bytes):
    print("going into extracting from pdf")
    pdf = fitz.open(
        stream=file_bytes,
        filetype="pdf"
    )

    text = ""

    for page in pdf:

        page_text = page.get_text("text") or ""

        text += page_text + "\n"

    if not text.strip():

        raise Exception(
            "No readable text extracted from PDF"
        )

    return text

def extract_text_from_docx_bytes(file_bytes):
    print("going into extract from docx")
    doc = Document(BytesIO(file_bytes))

    text = "\n".join(
        paragraph.text
        for paragraph in doc.paragraphs
    )

    if not text.strip():

        raise Exception(
            "No readable text extracted from DOCX"
        )

    return text

def extract_text_from_doc_bytes(file_bytes):

    import tempfile

    with tempfile.NamedTemporaryFile(
        suffix=".doc",
        delete=True
    ) as temp:

        temp.write(file_bytes)

        temp.flush()

        text = textract.process(
            temp.name
        ).decode("utf-8")

    if not text.strip():

        raise Exception(
            "No readable text extracted from DOC"
        )

    return text
def extract_text_from_image_bytes(file_bytes):

    image = Image.open(
        BytesIO(file_bytes)
    )

    text = pytesseract.image_to_string(
        image
    )

    if not text.strip():

        raise Exception(
            "No readable text extracted from image"
        )

    return text
def extract_resume_text(url):

    response = requests.get(url)

    if response.status_code != 200:
        raise Exception("Failed to fetch file")

    file_bytes = response.content

    filename = Path(
        urlparse(url).path
    ).name.lower()

    if filename.endswith(".pdf"):

        return extract_text_from_pdf_bytes(
            file_bytes
        )

    elif filename.endswith(".docx"):

        return extract_text_from_docx_bytes(
            file_bytes
        )

    elif filename.endswith(".doc"):

        return extract_text_from_doc_bytes(
            file_bytes
        )
    elif (
        filename.endswith(".png")
        or filename.endswith(".jpg")
        or filename.endswith(".jpeg")
    ):

        return extract_text_from_image_bytes(
            file_bytes
        )
    else:

        raise Exception(
            "Unsupported file format"
        )
# # -----------------------------
# # Helper: Extract text from PDF
# # -----------------------------
# def extract_text_from_pdf_url(url):
#     response = requests.get(url)
#     if response.status_code != 200:
#         raise Exception("Failed to fetch PDF from URL")
#     pdf_bytes = response.content
#     pdf = fitz.open(stream=pdf_bytes, filetype="pdf")
#     text = ""
#     for page in pdf:
#         page_text = page.get_text("text") or ""
#         text += page_text + "\n"
#     if not text.strip():
#         raise Exception("No readable text extracted from PDF")
#     return text
# def extract_text_from_docx_bytes(file_bytes):
#     doc = Document(BytesIO(file_bytes))

#     text = "\n".join(
#         paragraph.text
#         for paragraph in doc.paragraphs
#     )

#     return text
# -----------------------------
# Unified Skill Extraction Endpoint
# -----------------------------
@api_view(["POST"])
def generate_profile(request):
    user_id = request.data.get("user_id")
    resume_url = request.data.get("resume_url")
    manual_skills_list = request.data.get("skills", [])
    github_username = request.data.get("github_username")

    if not user_id:
        return Response({"error": "user_id is required"}, status=400)

    all_skills = set()

    # 1️⃣ Resume extraction
    if resume_url:
        try:
            text = extract_resume_text(
                resume_url
            )
            if not is_cs_resume(text):
                return Response(
                    {
                        "error": "Invalid resume. Please upload a Computer Science or IT-related resume."
                    },
                    status=400
                )
            resume_skills = extract_skills(text)
            all_skills.update(resume_skills)

            # Save/update resume record
            Resume.objects.update_or_create(
                user_id=user_id,
                defaults={"resume_url": resume_url, "parsed_text": text},
            )
        except Exception as e:
            return Response({"error": f"Resume processing failed: {str(e)}"}, status=500)

    # 2️⃣ Manual skills
    if manual_skills_list:
        all_skills.update([s.strip() for s in manual_skills_list if s.strip()])

    # 3️⃣ GitHub extraction
    if github_username:
        try:
            repos_url = f"https://api.github.com/users/{github_username}/repos"
            repos = requests.get(repos_url).json()
            if isinstance(repos, dict) and repos.get("message") == "Not Found":
                return Response({"error": "GitHub user not found"}, status=404)
            
            github_skills = set()
            for repo in repos:
                if repo.get("language"):
                    github_skills.add(repo["language"].lower())
                for topic in repo.get("topics", []):
                    github_skills.add(topic.lower())
            all_skills.update(github_skills)
        except Exception as e:
            return Response({"error": f"GitHub processing failed: {str(e)}"}, status=500)

    # 4️⃣ Save skills to DB without duplicates
    existing_skills = set(Skill.objects.filter(user_id=user_id).values_list("name", flat=True))
    new_skills = [s for s in all_skills if s not in existing_skills]

    for i, skill in enumerate(new_skills):
        Skill.objects.create(
            user_id=user_id,
            name=skill,
            level=50 + i * 5 if i * 5 <= 50 else 95,
            confidence=round(0.6 + (i / len(new_skills)) * 0.4, 2) if new_skills else 0.8
        )

    total_skills = Skill.objects.filter(user_id=user_id)
    score = min(total_skills.count() * 10, 100)

    return Response({
        "message": "Profile generated successfully",
        "skills": [{"name": s.name, "level": s.level, "confidence": s.confidence} for s in total_skills],
        "score": score
    })
